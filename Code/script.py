import fitz  # PyMuPDF
import os
import openai  # For GPT-4 (replace with Hugging Face if desired)
from scholarly import scholarly
import PyPDF2
from docx import Document
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
import time
import random
import glob
from pyzotero import zotero
import pyautogui
import time
import win32com.client as win32
import subprocess
import json


# Initialize OpenAI API key for GPT (optional if using Hugging Face)
openai.api_key = ''


# Function to extract text from the first page of PDFs
def extract_first_page_from_pdf(file_path):
    doc = fitz.open(file_path)
    first_page = doc.load_page(0)  # Load the first page
    text = first_page.get_text("text")
    return text

# Function to extract text from a PDF
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text()
        return text


def extract_first_page_from_word(doc_path, num_paragraphs=5):
    """
    Extracts the first few paragraphs from a Word document to approximate the first page content.

    Args:
        doc_path (str): Path to the Word document.
        num_paragraphs (int): Number of paragraphs to extract from the beginning of the document.

    Returns:
        str: Text content from the first few paragraphs.
    """
    doc = Document(doc_path)
    content = []

    # Collect the first `num_paragraphs` paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        if i < num_paragraphs:
            text = paragraph.text.strip()
            if text:  # Only include non-empty paragraphs
                content.append(text)
        else:
            break

    # Join the paragraphs into a single string
    return "\n".join(content)

# Function to predict the title using GPT-4 with the new API call
def predict_title(text):
    response = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": f"From the following text, strictly reply to me with only the title of the document:\n\n{text}",
                "temperature": 0.2,
            }
        ]
    )
    return response.choices[0].message.content

# Main function to process multiple PDFs and extract titles
def extract_titles_from_pdfs(folder_path):
    titles = []
    for pdf_file in os.listdir(folder_path):
        if pdf_file.endswith(".pdf") and not pdf_file.startswith("~$"):
            try:
                file_path = os.path.join(folder_path, pdf_file)
                text = extract_first_page_from_pdf(file_path)
                title = predict_title(text)
                titles.append([title,file_path])
            except Exception as e:
                print(f"Error processing {pdf_file}: {e}")
    return titles



def summarize_paper(research_title, text, previous_summary="", max_tokens=300):
    if previous_summary == "":
        prompt = f"Write a focused literature review for the research titled '{research_title}',based strictly on the following text. Concentrate on synthesizing findings and main ideas without general introductions or descriptions of what a literature review entails.Exclude mentions of specific author names or years. Here is the reference text: {text}"
    response = openai.chat.completions.create(
        model="gpt-4",
        max_tokens=max_tokens,
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )
    return response.choices[0].message.content


def run_uipath_workflow(workflow_path, citation_paper_title, document_title):
    # Prepare the input arguments in JSON format
    input_args = {
        'citation_title': citation_paper_title,
        'document_title': document_title
    }
    input_args_json = json.dumps(input_args)



    # Command to run the UiPath workflow
    command = [
        r"C:\Users\xsilv\AppData\Local\Programs\UiPath\Studio\UiRobot.exe",  # Path to UiPath Robot executable
        '-file', workflow_path,  # Path to UiPath workflow (.xaml)
        '--input', input_args_json  # Pass input arguments as JSON
    ]

    #print(input_args_json)
    #print(command)
    # Run the command
    result = subprocess.run(command, capture_output=True, text=True)

    # Output the result
    #print(result.stdout)
    #print(result.stderr)

# Function to find the start and end of the EndNote Bibliography section
def find_reference_section(doc):
    for i, para in enumerate(doc.paragraphs):
        # Check if it's the EndNote Bibliography style
        if para.style.name == "EndNote Bibliography Title":
            return i
        elif para.style.name == "EndNote Bibliography":
            return i
    return None  # No reference section found

#Function to save the summary to a Word document
def save_summary_to_word(summary, doc_name="summary.docx"):
    print("Writing summary...")
    if not os.path.exists(doc_name):
        doc = Document()
        doc.add_heading('Summary', 0)
    else:
        doc = Document(doc_name)

    # Find the reference section
    ref_index = find_reference_section(doc)

    # If reference section exists, move the new content just above the reference section
    if ref_index is not None:
        # Move the new paragraph above the reference section
        ref_paragraph = doc.paragraphs[ref_index]
        # Insert a new paragraph before the reference paragraph
        new_paragraph = ref_paragraph.insert_paragraph_before(summary)
        new_paragraph.add_run(" [MISSING CITATION]")
    else:
        # If no reference section exists, content stays at the end
        doc.add_paragraph(summary).add_run(" [MISSING CITATION]")

    # Save the document
    doc.save(doc_name)
    print(f"Summary saved to {doc_name}")

# Main function to process the PDF
def process_pdf(pdf_path,doc_name, research_title):
    text = extract_text_from_pdf(pdf_path)

    # Split the text if it's too long
    max_chunk_size = 3000  # Adjust as needed for chunking
    if len(text) > max_chunk_size:
        chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
    else:
        chunks = [text]

    #print(chunks)
    #full_summary = ""
    #for chunk in chunks:
    #    summary = summarize_paper(chunk)
    #    full_summary += summary + "\n\n"

    #testing
    full_summary = summarize_paper(research_title, chunks[0])
    save_summary_to_word(full_summary,doc_name)


# Example usage

def main():
    #CHANGE FILE PATH ACCORDINGLY
    base_path = r'C:\Users\xsilv\OneDrive\Desktop\School work\FYPJ'
    folder_path = "C:/Users/xsilv/OneDrive/Desktop/School work/FYPJ/pdf"
    workflow_path = r'C:\Users\xsilv\OneDrive\Desktop\School work\FYPJ\EndNote.1.0.9.nupkg'
    titles = extract_titles_from_pdfs(folder_path)

    #print(titles)

    #finding research paper filename

    doc_name = ""
    while not os.path.exists(doc_name):
        doc_name = input("Input an existing document filename: ")
        if not doc_name.endswith(".docx"):
            doc_name += ".docx"

    print("Document is found!")
    #recognise the title in the research paper (word doc)
    word_first_page = extract_first_page_from_word(doc_name,10)
    #print(word_first_page)
    research_title = predict_title(word_first_page)
    print("Research title is:")
    print(research_title)

    doc_location = os.path.join(base_path,doc_name)
    #input to chatgpt the title of the paper and ask to turn the paper into a literature review
    for title,pdf_path in titles:
        process_pdf(pdf_path,doc_name,research_title)
        run_uipath_workflow(workflow_path, title, doc_location)


# Execute the move function
main()